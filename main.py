import tkinter as tk
from tkinter import messagebox, simpledialog, Toplevel
from tkinter import ttk
from pathlib import Path
from datetime import date
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from deep_translator import GoogleTranslator
import os
import webbrowser
import ast
import requests

# Ruta base
BASE_DIR = Path(__file__).resolve().parent
ESP_FILE = BASE_DIR / 'esp.docx'
ENG_FILE = BASE_DIR / 'eng.docx'
LISTADO_FILE = BASE_DIR / 'listado.docx'
BBDD_FILE = BASE_DIR / 'sugerencias_bbdd.docx'
LOGO_PATH = BASE_DIR / 'logo.jpeg'
ICONO_PATH = BASE_DIR / 'logo.ico'

# Función para cargar sugerencias desde Word
def cargar_sugerencias():
    sugerencias = {}
    if BBDD_FILE.exists():
        doc = Document(BBDD_FILE)
        for p in doc.paragraphs:
            if ':' in p.text:
                try:
                    clave, datos = p.text.split("::", 1) if '::' in p.text else p.text.split(":", 1)
                    clave = int(clave.strip())
                    valores = ast.literal_eval(datos.strip())
                    if isinstance(valores, list) and len(valores) == 3:
                        sugerencias[clave] = valores
                except:
                    continue
    return sugerencias

# Guardar sugerencias en Word
def guardar_bbdd(sugerencias):
    doc = Document()
    for clave, datos in sorted(sugerencias.items()):
        doc.add_paragraph(f"{clave}: {datos}")
    doc.save(BBDD_FILE)

# Traducción
DEEPL_API_KEY = "1e2d1343-5e72-4dd7-b6ae-7bbb93f30c96:fx"

def traducir_a_ingles(texto_es):
    url = "https://api-free.deepl.com/v2/translate"
    headers = {
        "Content-Type": "application/x-www-form-urlencoded"
    }
    data = {
        "auth_key": DEEPL_API_KEY,
        "text": texto_es,
        "target_lang": "EN"
    }
    try:
        response = requests.post(url, headers=headers, data=data)
        response.raise_for_status()
        result = response.json()
        return result["translations"][0]["text"]
    except Exception as e:
        print("Error al traducir:", e)
        print("Respuesta del servidor:", response.text)
        return texto_es + " (no traducido)"

# Crear sugerencia nueva
def crear_sugerencia():
    sugerencias = cargar_sugerencias()
    ventana = Toplevel()
    ventana.title("Nueva sugerencia")
    if ICONO_PATH.exists(): ventana.iconbitmap(ICONO_PATH)

    frame = tk.Frame(ventana, padx=10, pady=10)
    frame.pack()

    tk.Label(frame, text="Precio (€):").grid(row=0, column=0, sticky="e", pady=5)
    entry_precio = tk.Entry(frame)
    entry_precio.grid(row=0, column=1, pady=5)

    tk.Label(frame, text="Descripción en español:").grid(row=1, column=0, sticky="ne", pady=5)
    text_desc = tk.Text(frame, height=4, width=40)
    text_desc.grid(row=1, column=1, pady=5)

    def guardar():
        try:
            precio_str = entry_precio.get().replace(",", ".")
            precio = float(precio_str)
            desc_es = text_desc.get("1.0", tk.END).strip()
            desc_en = traducir_a_ingles(desc_es)
            nueva_clave = max(sugerencias.keys(), default=0) + 1
            sugerencias[nueva_clave] = [precio, desc_es, desc_en]
            guardar_bbdd(sugerencias)
            messagebox.showinfo("Éxito", f"Sugerencia {nueva_clave} añadida.", parent=ventana)
            ventana.destroy()
        except Exception as e:
            messagebox.showerror("Error", f"Datos inválidos.\n\nDetalles: {str(e)}", parent=ventana)

    tk.Button(frame, text="Guardar sugerencia", command=guardar).grid(row=2, columnspan=2, pady=10)

# Modificar sugerencia existente
def modificar_sugerencia():
    sugerencias = cargar_sugerencias()
    clave = simpledialog.askinteger("Modificar sugerencia", "Número de sugerencia a modificar:")
    if not clave or clave not in sugerencias:
        messagebox.showerror("Error", "Número inválido o no existente")
        return

    ventana = Toplevel()
    ventana.title(f"Modificar sugerencia {clave}")
    if ICONO_PATH.exists(): ventana.iconbitmap(ICONO_PATH)

    frame = tk.Frame(ventana, padx=10, pady=10)
    frame.pack()

    tk.Label(frame, text="Precio (€):").grid(row=0, column=0, sticky="e", pady=5)
    entry_precio = tk.Entry(frame)
    entry_precio.insert(0, sugerencias[clave][0])
    entry_precio.grid(row=0, column=1, pady=5)

    tk.Label(frame, text="Descripción en español:").grid(row=1, column=0, sticky="ne", pady=5)
    text_desc = tk.Text(frame, height=4, width=40)
    text_desc.insert("1.0", sugerencias[clave][1])
    text_desc.grid(row=1, column=1, pady=5)

    def guardar():
        try:
            nuevo_precio = float(entry_precio.get().replace(",", "."))
            nueva_desc_es = text_desc.get("1.0", tk.END).strip()
            nueva_desc_en = traducir_a_ingles(nueva_desc_es)
            sugerencias[clave] = [nuevo_precio, nueva_desc_es, nueva_desc_en]
            guardar_bbdd(sugerencias)
            messagebox.showinfo("Modificado", f"Sugerencia {clave} modificada con éxito")
            ventana.destroy()
        except:
            messagebox.showerror("Error", "Datos inválidos", parent=ventana)

    tk.Button(frame, text="Guardar cambios", command=guardar).grid(row=2, columnspan=2, pady=10)

# Crear listado Word
def crear_listado():
    sugerencias = cargar_sugerencias()
    if not sugerencias:
        messagebox.showwarning("Aviso", "No hay sugerencias guardadas")
        return

    doc = Document()
    doc.add_heading("Listado de Sugerencias", level=1)

    for clave in sorted(sugerencias):
        precio, desc_es, desc_en = sugerencias[clave]
        p = doc.add_paragraph()
        p.add_run(f"{clave}. ").bold = True
        p.add_run(f"{desc_es} ({precio}€)")

    doc.save(LISTADO_FILE)
    messagebox.showinfo("Listado generado", f"Listado guardado en {LISTADO_FILE.name}")

# Crear menú Word
from pathlib import Path
from docx.shared import Cm
BASE_DIR = Path(__file__).resolve().parent
ESP_FILE = BASE_DIR / "esp.docx"
ENG_FILE = BASE_DIR / "eng.docx"

def crear_menu():
    sugerencias = cargar_sugerencias()
    seleccion = simpledialog.askstring("Menú", "Números de sugerencia (ej: 1,2,3):")
    if not seleccion:
        return

    if any('.' in s for s in seleccion.split(" ")):
        messagebox.showerror("Error", "Formato incorrecto. Por favor, separa los números con comas, no con puntos.")
        return

    try:
        indices = [int(i.strip()) for i in seleccion.replace(".", ",").split(",") if i.strip().isdigit() and int(i.strip()) in sugerencias]
    except Exception as e:
        messagebox.showerror("Error", f"Entrada inválida: {str(e)}\nAsegúrate de separar los números con comas.")
        return

    hoy = date.today().isoformat()
    output_dir = BASE_DIR / 'menus'
    output_dir.mkdir(exist_ok=True)

    plantilla_esp = Document(ESP_FILE)
    plantilla_eng = Document(ENG_FILE)

    def agregar_sugerencias(doc, idioma):
        for duplicado in range(2):  # Dos bloques
            if duplicado == 1:
                doc.add_paragraph()  # Separación entre copias
                doc.add_paragraph()  # Separación entre copias

            for idx, num in enumerate(indices):
                precio, desc_es, desc_en = sugerencias[num]
                texto = desc_es if idioma == 'ES' else desc_en
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(2.5)
                p.paragraph_format.right_indent = Cm(2.5)

                run_texto = p.add_run(f"- {texto} - ")
                run_precio = p.add_run(f"{precio}€")
                for r in [run_texto, run_precio]:
                    r.font.name = 'Georgia'
                    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Georgia')
                    r.font.size = Pt(12)
                color = RGBColor(0, 0, 0) if idx % 2 == 0 else RGBColor(0, 0, 255)
                run_texto.font.color.rgb = color
                run_precio.font.color.rgb = color
                run_precio.bold = True
                p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # Título al final del bloque
            titulo = doc.add_paragraph()
            titulo.paragraph_format.left_indent = Cm(0)
            titulo.paragraph_format.right_indent = Cm(0)
            texto_titulo = "SUGERENCIAS" if idioma == 'ES' else "SUGGESTIONS"
            run_titulo = titulo.add_run(texto_titulo)
            run_titulo.bold = True
            run_titulo.font.size = Pt(14)
            run_titulo.font.color.rgb = RGBColor(0, 0, 255)
            titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    agregar_sugerencias(plantilla_esp, 'ES')
    agregar_sugerencias(plantilla_eng, 'EN')

    plantilla_esp.save(output_dir / f"esp_{hoy}.docx")
    plantilla_eng.save(output_dir / f"eng_{hoy}.docx")
    messagebox.showinfo("Menú creado", "Menús Word generados correctamente")
# Interfaz principal
def interfaz():
    root = tk.Tk()
    root.title("Sugerchef")
    if ICONO_PATH.exists(): root.iconbitmap(ICONO_PATH)
    root.configure(bg="#E6F2FA")

    if LOGO_PATH.exists():
        from PIL import ImageTk, Image
        img = Image.open(LOGO_PATH)
        img = img.resize((120, 120), Image.Resampling.LANCZOS)
        logo = ImageTk.PhotoImage(img)
        tk.Label(root, image=logo, bg="#E6F2FA").pack(pady=(10, 0))

    tk.Label(root, text="Sugerencias para Menú", font=("Georgia", 14, "bold"), bg="#E6F2FA", fg="#003E6B").pack(pady=5)

    btn_frame = tk.Frame(root, bg="#E6F2FA")
    btn_frame.pack(pady=10)

    tk.Button(btn_frame, text="Crear sugerencia", command=crear_sugerencia, width=25).grid(row=0, column=0, padx=5, pady=5)
    tk.Button(btn_frame, text="Modificar sugerencia", command=modificar_sugerencia, width=25).grid(row=0, column=1, padx=5, pady=5)
    tk.Button(btn_frame, text="Crear menú Word", command=crear_menu, width=25).grid(row=1, column=0, padx=5, pady=5)
    tk.Button(btn_frame, text="Imprimir listado completo", command=crear_listado, width=25).grid(row=1, column=1, padx=5, pady=5)

    info_frame = tk.Frame(root, bg="#E6F2FA")
    info_frame.pack(pady=10)

    autor = tk.Label(info_frame, text="App creada por Francisco Javier Fiestas Botella", font=("Georgia", 9), bg="#E6F2FA", fg="#003E6B")
    autor.pack()

    github = tk.Label(info_frame, text="GitHub: Javierfiestasbotella", font=("Georgia", 9, "underline"), fg="blue", cursor="hand2", bg="#E6F2FA")
    github.pack()
    github.bind("<Button-1>", lambda e: webbrowser.open("https://github.com/Javierfiestasbotella"))

    video = tk.Label(info_frame, text="Ver explicación en YouTube", font=("Georgia", 9, "underline"), fg="blue", cursor="hand2", bg="#E6F2FA")
    video.pack()
    video.bind("<Button-1>", lambda e: webbrowser.open("https://youtu.be/7ejvejTGhWk"))

    telefono = tk.Label(info_frame, text="Contacto: 628796613", font=("Georgia", 9), bg="#E6F2FA", fg="#003E6B")
    telefono.pack()

    root.mainloop()

if __name__ == "__main__":
    interfaz()
